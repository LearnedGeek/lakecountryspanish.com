"""
Converts the hand-authored lesson proposals (Markdown + YAML frontmatter)
under `docs/lcs-business/curriculum/proposals/*.md` into template-shaped
.docx files the DocxLessonParser can ingest.

The proposals already have the structured data we need (vocab tiered,
songs with URLs, standards with WI codes); this script just reshapes it
into the Word table layout the parser expects.

Output drops alongside the source .md files as `*.template.docx`.

Skips construyendo-los-cimientos-5to-grado.md — it's a multi-chapter
grade orientation, not a single lesson, so it doesn't fit the per-lesson
template shape. (Different artifact, separate template later if needed.)
"""

import re
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    print("ERROR: pyyaml not installed. Run: pip install pyyaml", file=sys.stderr)
    sys.exit(1)

from docx import Document
from docx.shared import Pt


PROPOSALS_DIR = Path(__file__).resolve().parent.parent / "docs" / "lcs-business" / "curriculum" / "proposals"
SKIP = {"construyendo-los-cimientos-5to-grado.md"}

# Body section heading names (the markdown bodies use these as ## headers).
# Maps from the proposal heading text → the template Heading 1 name.
SECTION_HEADING_MAP = {
    "Apertura (Opening)":                              "Apertura",
    "Apertura":                                        "Apertura",
    "Presentación (Presentation)":                     "Presentación",
    "Presentación":                                    "Presentación",
    "Extensión para 1.º y 2.º grado":                  "Extensión",
    "Extensión":                                       "Extensión",
    "Actividad: Colorear":                             "Actividad",
    "Actividad":                                       "Actividad",
    "Opciones de juegos":                              "Juegos",
    "Opciones de Juegos":                              "Juegos",
    "Cultura":                                         "Cultura",
    "Cierre (Closing)":                                "Cierre",
    "Cierre":                                          "Cierre",
}

ROLE_MAP = {
    "opening": "Opening",
    "presentation": "Presentation",
    "closing": "Closing",
}


def parse_frontmatter_and_body(md_text):
    """Returns (frontmatter_dict, body_after_frontmatter)."""
    if not md_text.startswith("---"):
        raise ValueError("Markdown missing YAML frontmatter")
    end = md_text.index("\n---", 4)
    fm_text = md_text[4:end]
    body = md_text[end + 4:].lstrip("\n")
    return yaml.safe_load(fm_text), body


def grade_band_text(value):
    if not value:
        return "K-2"
    return str(value).replace("–", "-")


def sessions_int(fm):
    """Pull a session count integer. Accepts 'sessions: "3 clases"' or 'day: 1'."""
    if "sessions" in fm and fm["sessions"]:
        m = re.search(r"\d+", str(fm["sessions"]))
        if m:
            # If a range like "2-3", take the higher end (covers worst case).
            nums = [int(n) for n in re.findall(r"\d+", str(fm["sessions"]))]
            return max(nums)
    return 1


def duration_int(fm):
    return int(fm.get("duration_minutes_per_session") or fm.get("duration_minutes") or 30)


def trimester_int(fm):
    return int(fm.get("trimester") or 0)


def split_body_sections(body):
    """Split the markdown body into a dict[heading_name → content]."""
    out = {}
    current_heading = None
    current_lines = []

    def flush():
        if current_heading is None:
            return
        mapped = SECTION_HEADING_MAP.get(current_heading.strip())
        if mapped is None:
            return
        text = "\n".join(current_lines).strip()
        # Strip blockquote markers and tidy bullet artifacts so the prose
        # looks natural inside the .docx paragraphs.
        text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
        out[mapped] = text

    for line in body.splitlines():
        if line.startswith("## "):
            flush()
            current_heading = line[3:].strip()
            current_lines = []
        elif line.startswith("# "):
            # top-level heading; ignore
            flush()
            current_heading = None
            current_lines = []
        else:
            if current_heading is not None:
                current_lines.append(line)
    flush()
    return out


# ----------- docx writers (mirror the structure in make_lesson_template.py) -----------

def add_heading1(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 1"]
    return p


def make_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v) if v is not None else ""
    return table


def make_tier_table(doc, tiers):
    rows = sum(len(words) for _, words in tiers) + 1
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Word"
    r = 1
    for tier, words in tiers:
        for w in words:
            table.cell(r, 0).text = tier
            table.cell(r, 1).text = str(w)
            r += 1
    return table


def make_songs_table(doc, songs):
    table = doc.add_table(rows=len(songs) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Title"
    table.cell(0, 1).text = "Role"
    table.cell(0, 2).text = "URL"
    for i, song in enumerate(songs, start=1):
        table.cell(i, 0).text = song.get("title", "")
        role_raw = song.get("role", "supplemental")
        table.cell(i, 1).text = ROLE_MAP.get(role_raw.lower(), "Supplemental")
        table.cell(i, 2).text = song.get("link", "") or ""
    return table


def make_standards_table(doc, standards):
    rows = [(s.get("wi_code", ""), s.get("addressed_by", "")) for s in standards]
    if not rows:
        return None
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "WI Code"
    table.cell(0, 1).text = "Notes"
    for i, (code, notes) in enumerate(rows, start=1):
        table.cell(i, 0).text = code
        table.cell(i, 1).text = notes
    return table


def write_docx(fm, sections, out_path):
    doc = Document()

    # Metadata
    add_heading1(doc, "Metadata")
    make_kv_table(doc, [
        ("Title", fm.get("lesson", "")),
        ("Unit", fm.get("unit", "")),
        ("Grade Band", grade_band_text(fm.get("grade_band"))),
        ("Sessions", sessions_int(fm)),
        ("Trimester", trimester_int(fm)),
        ("Theme", fm.get("theme", "")),
        ("Objective", fm.get("objective", "")),
        ("Duration Minutes", duration_int(fm)),
    ])

    # Vocabulary
    add_heading1(doc, "Vocabulary")
    vocab = fm.get("vocabulary") or {}
    tiers = []
    if vocab.get("core"):     tiers.append(("Core", vocab["core"]))
    if vocab.get("stretch"):  tiers.append(("Stretch", vocab["stretch"]))
    if vocab.get("challenge"): tiers.append(("Challenge", vocab["challenge"]))
    if tiers:
        make_tier_table(doc, tiers)

    # Materials
    add_heading1(doc, "Materials")
    for m in fm.get("materials") or []:
        doc.add_paragraph(str(m))

    # Songs
    add_heading1(doc, "Songs")
    songs = fm.get("songs") or []
    if songs:
        # Drop TODO links so the cell is empty — parser skips rows missing URL.
        cleaned = []
        for s in songs:
            link = s.get("link", "")
            if not link or link.strip().upper() == "TODO":
                link = ""
            cleaned.append({**s, "link": link})
        # Only emit rows that have a real URL — parser requires URL.
        with_urls = [s for s in cleaned if s.get("link")]
        if with_urls:
            make_songs_table(doc, with_urls)

    # Body sections in template order
    for section in ("Apertura", "Presentación", "Extensión",
                    "Actividad", "Juegos", "Cultura", "Cierre"):
        text = sections.get(section)
        if not text:
            continue
        add_heading1(doc, section)
        for paragraph in text.split("\n\n"):
            cleaned = paragraph.strip()
            if cleaned:
                doc.add_paragraph(cleaned)

    # Standards
    standards = fm.get("standards") or []
    if standards:
        add_heading1(doc, "Standards")
        make_standards_table(doc, standards)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    if not PROPOSALS_DIR.exists():
        print(f"ERROR: proposals directory not found: {PROPOSALS_DIR}", file=sys.stderr)
        sys.exit(1)

    converted = []
    for md_path in sorted(PROPOSALS_DIR.glob("*.md")):
        if md_path.name in SKIP:
            print(f"skip: {md_path.name} (not a single-lesson artifact)")
            continue

        try:
            fm, body = parse_frontmatter_and_body(md_path.read_text(encoding="utf-8"))
            sections = split_body_sections(body)
            out_path = md_path.with_suffix(".template.docx")
            write_docx(fm, sections, out_path)
            converted.append(out_path)
            print(f"wrote: {out_path.name}  ({out_path.stat().st_size:,} bytes)")
        except Exception as e:
            print(f"FAIL  {md_path.name}: {e}", file=sys.stderr)

    print(f"\n{len(converted)} file(s) converted.")
    print("Upload each via /Curriculum/Upload to test the parser end-to-end.")


if __name__ == "__main__":
    main()
