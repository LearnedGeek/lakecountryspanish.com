"""
Generates the LCS lesson template .docx that Karen + Cece download to author
new lessons against. Output drops into:

    src/LakeCountrySpanish.Web/wwwroot/curriculum/lesson-template.docx

The parser (Services/Curriculum/DocxLessonParser.cs) reads documents in
exactly this shape. Karen edits the cells/sections without touching the
table structure; the parser maps cells → fields by position.

Also generates a SAMPLE (filled with the La Familia content) at:

    src/LakeCountrySpanish.Web/wwwroot/curriculum/sample-la-familia.docx

This sample doubles as a smoke-test fixture: upload it through
/Curriculum/Upload to confirm the pipeline produces a Day + LessonVideos.
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def add_heading1(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Heading 1"]
    return p


def make_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    return table


def make_tier_table(doc, tiers):
    rows = sum(len(words) for _, words in tiers) + 1  # +1 header
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Word"
    r = 1
    for tier, words in tiers:
        for w in words:
            table.cell(r, 0).text = tier
            table.cell(r, 1).text = w
            r += 1
    return table


def make_songs_table(doc, songs):
    table = doc.add_table(rows=len(songs) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "Title"
    table.cell(0, 1).text = "Role"
    table.cell(0, 2).text = "URL"
    for i, (t, r, u) in enumerate(songs, start=1):
        table.cell(i, 0).text = t
        table.cell(i, 1).text = r
        table.cell(i, 2).text = u
    return table


def make_standards_table(doc, codes):
    table = doc.add_table(rows=len(codes) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "WI Code"
    table.cell(0, 1).text = "Notes (optional)"
    for i, code in enumerate(codes, start=1):
        table.cell(i, 0).text = code
    return table


def make_blank_template(out_path: Path):
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Lake Country Spanish — Lesson Template")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph(
        "Fill in this template and upload via /Curriculum/Upload. "
        "Don't change the table structures — the parser maps cells by position. "
        "Heading 1 styled section names are how the parser finds each section."
    )

    add_heading1(doc, "Metadata")
    make_kv_table(doc, [
        ("Title", "[Lesson title]"),
        ("Unit", "[Existing Unit name]"),
        ("Grade Band", "K-2 / 3-5 / 6-8 (or specific: K5, Grade3, etc.)"),
        ("Sessions", "1"),
        ("Trimester", "1"),
        ("Theme", "[short-slug]"),
        ("Objective", "[One-sentence learning objective]"),
        ("Duration Minutes", "30"),
    ])

    add_heading1(doc, "Vocabulary")
    doc.add_paragraph(
        "List vocab terms tiered by difficulty. Core = everyone learns these. "
        "Stretch = covered as the group is ready. Challenge = advanced."
    )
    make_tier_table(doc, [
        ("Core", ["word1", "word2", "word3"]),
        ("Stretch", ["word4", "word5"]),
        ("Challenge", ["word6"]),
    ])

    add_heading1(doc, "Materials")
    doc.add_paragraph("Material 1")
    doc.add_paragraph("Material 2")
    doc.add_paragraph("Material 3")

    add_heading1(doc, "Songs")
    doc.add_paragraph(
        "Each song row becomes a LessonVideo. Role values: "
        "Opening / Presentation / Closing / Supplemental."
    )
    make_songs_table(doc, [
        ("Buenos días", "Opening", "https://youtu.be/..."),
        ("[Vocab song]", "Presentation", "https://youtu.be/..."),
        ("Adiós, amigos", "Closing", "https://youtu.be/..."),
    ])

    add_heading1(doc, "Apertura")
    doc.add_paragraph("Comenzamos la clase con la canción de saludo. …")

    add_heading1(doc, "Presentación")
    doc.add_paragraph("Uso de tarjetas / repetición / canción de presentación.")

    add_heading1(doc, "Extensión")
    doc.add_paragraph("Práctica de oraciones / extensión para 1° y 2° grado.")

    add_heading1(doc, "Actividad")
    doc.add_paragraph("Hoja de colorear / actividad de la clase.")

    add_heading1(doc, "Juegos")
    doc.add_paragraph("1. [Juego 1]")
    doc.add_paragraph("2. [Juego 2]")
    doc.add_paragraph("3. [Juego 3]")

    add_heading1(doc, "Cultura")
    doc.add_paragraph("Bandera de un país hispanohablante — coloreamos.")

    add_heading1(doc, "Cierre")
    doc.add_paragraph("Terminamos la clase con la canción de despedida.")

    add_heading1(doc, "Standards")
    doc.add_paragraph(
        "WI codes the lesson aligns to. Pick from the seeded WI DPI Novice-band "
        "standards. The parser attaches by code lookup."
    )
    make_standards_table(doc, [
        "WL.IT.1.c.n1",
        "WL.IP.2.a.n1",
        "WL.PS.3.a.n1",
        "WL.IC.4.a.n+",
    ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote: {out_path}  ({out_path.stat().st_size:,} bytes)")


def make_la_familia_sample(out_path: Path):
    """A filled-in template using Karen's La Familia content."""
    doc = Document()
    add_heading1(doc, "Metadata")
    make_kv_table(doc, [
        ("Title", "La Familia"),
        ("Unit", "La Familia"),
        ("Grade Band", "K-2"),
        ("Sessions", "3"),
        ("Trimester", "1"),
        ("Theme", "familia"),
        ("Objective", "Los estudiantes podrán identificar los miembros de la familia."),
        ("Duration Minutes", "30"),
    ])

    add_heading1(doc, "Vocabulary")
    make_tier_table(doc, [
        ("Core", ["mamá", "papá", "hermano", "hermana", "perro", "gato", "mascota"]),
        ("Stretch", ["abuelo", "abuela", "bebé", "tío", "tía", "primo", "prima"]),
    ])

    add_heading1(doc, "Materials")
    for m in [
        "Tarjetas con fotos de cada miembro de la familia",
        "Hoja de trazado de palabras",
        "Hoja de match (familia)",
        "Libro de la familia para colorear",
        "Plantilla de bingo de la familia",
        "Pelota",
        "Sillas para 'Las sillas musicales'",
        "Bandera de Chile (u otro país hispanohablante)",
    ]:
        doc.add_paragraph(m)

    add_heading1(doc, "Songs")
    make_songs_table(doc, [
        ("Buenos días, ¿cómo estás?", "Opening", "https://youtu.be/X-UXGWTEfEI?si=QCQYSD8HJP-wIaqq"),
        ("Mi familia, mi familia", "Presentation", "https://youtu.be/Smm7csgMoMs?si=uT9Ai94GRKUfmGg9"),
        ("Adiós, amigos", "Closing", "https://youtu.be/..."),
    ])

    add_heading1(doc, "Apertura")
    doc.add_paragraph(
        "Comenzamos la clase con la canción del saludo. ¡Buenos días! "
        "¿Cómo estás? Muy bien, gracias…"
    )

    add_heading1(doc, "Presentación")
    doc.add_paragraph(
        "Uso de las tarjetas con cada miembro de la familia. Se repite varias "
        "veces el vocabulario junto a la clase. Puedes usar también la hoja "
        "de match y repetir con ellos cada vez que hacen match."
    )

    add_heading1(doc, "Extensión")
    doc.add_paragraph(
        "Practicar oraciones personales: Mi papá se llama ___. Mi mamá se "
        "llama ___. Mi hermana se llama ___. Mi perro se llama ___."
    )

    add_heading1(doc, "Actividad")
    doc.add_paragraph(
        "Coloreamos nuestro libro de la familia y trazamos las palabras."
    )

    add_heading1(doc, "Juegos")
    for g in [
        "1. Bingo de la familia — se usa la plantilla.",
        "2. La pelota de la familia — pasan la pelota mientras suena la música; "
        "cuando para, el niño dice un miembro o responde una pregunta como "
        "'¿Quién es la mamá de tu mamá?' → La abuela.",
        "3. Dibuja a tu familia — cada niño dibuja su familia y la presenta: "
        "'Mi papá es… Mi mamá se llama…'",
        "4. Memoria familiar — pares de tarjetas boca abajo; al encontrar un par "
        "dicen el nombre en español.",
        "5. Las sillas musicales — el que no encuentra asiento dice un miembro "
        "de la familia para seguir en el juego.",
    ]:
        doc.add_paragraph(g)

    add_heading1(doc, "Cultura")
    doc.add_paragraph(
        "Introducimos la bandera de un país hispanohablante y la coloreamos en "
        "clase. Sugerencia: Chile."
    )

    add_heading1(doc, "Cierre")
    doc.add_paragraph(
        "Terminamos la clase con la canción de despedida. ¡Adiós, amigos! "
        "Hasta luego. Hasta la vista. Chao, chao."
    )

    add_heading1(doc, "Standards")
    make_standards_table(doc, [
        "WL.IT.1.c.n1",
        "WL.IP.2.a.n1",
        "WL.PS.3.a.n1",
        "WL.IC.4.a.n+",
    ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote: {out_path}  ({out_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    target = root / "src" / "LakeCountrySpanish.Web" / "wwwroot" / "curriculum"
    make_blank_template(target / "lesson-template.docx")
    make_la_familia_sample(target / "sample-la-familia.docx")
